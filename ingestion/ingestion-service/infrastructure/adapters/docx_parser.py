import asyncio
from io import BytesIO
from uuid import uuid4
from typing import Tuple

from docx import Document as DocxDocument

from application.ports.i_document_parser import IDocumentParser
from domain.entities import Document
from domain.errors import EmptyDocumentError


class DocxParser(IDocumentParser):
    async def parse(self, content: bytes, filename: str, mime_type: str) -> Tuple[str, Document]:
        def _extract() -> str:
            doc = DocxDocument(BytesIO(content))
            return "\n".join(para.text for para in doc.paragraphs)

        text = await asyncio.to_thread(_extract)

        if not text or not text.strip():
            raise EmptyDocumentError(f"No text could be extracted from DOCX: {filename}")

        document = Document(
            id=str(uuid4()),
            filename=filename,
            mime_type=mime_type,
            content_source="upload",
        )
        return text, document
